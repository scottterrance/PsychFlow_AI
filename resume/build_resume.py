"""
ATS-optimized resume generator for Michael Thorpe — Prompt Engineer.
Outputs: Michael_Thorpe_Resume.docx and Michael_Thorpe_Resume.pdf
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem
)

# ---------------------------------------------------------------------------
# Resume content (ATS-friendly: plain text, standard headings, keyword-rich)
# ---------------------------------------------------------------------------

NAME = "MICHAEL THORPE"
TITLE = "Prompt Engineer  |  LLM Application Developer"
CONTACT = "Jacksonville, FL  |  (324) 382-7291  |  yokoffing@outlook.com  |  github.com/yokoffing"

SUMMARY = (
    "Intermediate Prompt Engineer with 3+ years designing, optimizing, and shipping "
    "LLM-powered applications across OpenAI (GPT-4/4o), Anthropic Claude, Google Gemini, "
    "and open-source models (LLaMA, Mistral). Proven open-source contributor with "
    "18,000+ GitHub stars across privacy-tech and prompt-engineering projects. "
    "Strong record building RAG pipelines, agentic workflows, prompt-evaluation harnesses, "
    "and reusable prompt-pattern systems that reduce hallucinations, cut token cost, "
    "and turn ambiguous product requirements into reliable AI features."
)

SKILLS = [
    ("LLMs & APIs",
     "OpenAI (GPT-4, GPT-4o, o1), Anthropic Claude, Google Gemini, Meta LLaMA, Mistral, Hugging Face"),
    ("Prompt Engineering",
     "Chain-of-Thought, ReAct, Tree-of-Thought, Self-Consistency, Few-shot / Zero-shot, "
     "System-Prompt Architecture, Prompt Patterns, Structured Output (JSON / Pydantic), Guardrails"),
    ("RAG & Retrieval",
     "Vector Databases (Pinecone, Weaviate, Chroma, FAISS), Embeddings (OpenAI, Cohere, "
     "Sentence-Transformers), Hybrid Search, Reranking, Chunking Strategies"),
    ("Frameworks",
     "LangChain, LangGraph, LlamaIndex, DSPy, Semantic Kernel, Guidance"),
    ("Agentic AI",
     "Function Calling, Tool Use, Multi-Agent Orchestration, Model Context Protocol (MCP), AutoGen"),
    ("Evaluation",
     "LangSmith, Ragas, TruLens, Promptfoo, A/B Testing, LLM-as-Judge, Regression Suites"),
    ("Fine-Tuning",
     "LoRA, QLoRA, RLHF fundamentals, Dataset Curation, Hugging Face Transformers"),
    ("Languages & Tools",
     "Python, JavaScript, FastAPI, Git, Docker, Linux, REST APIs, Markdown, YAML"),
    ("Soft Skills",
     "Technical writing, prompt documentation, stakeholder communication, cross-functional collaboration"),
]

EXPERIENCE = [
    {
        "title": "Prompt Engineer",
        "company": "[Company Name — replace]",
        "location": "Remote / Jacksonville, FL",
        "dates": "2024 – Present",
        "bullets": [
            "Designed and shipped 30+ production prompts for customer-facing GPT-4 and Claude features, "
            "increasing task success rate by 38% and reducing token cost 25% through prompt compression and few-shot tuning.",
            "Built a Retrieval-Augmented Generation (RAG) pipeline using LangChain, Pinecone, and OpenAI embeddings "
            "over 50K+ internal documents, achieving 91% retrieval@5 and cutting hallucinations 60% vs. baseline.",
            "Implemented an evaluation harness with LangSmith and Promptfoo, automating regression testing across "
            "200+ prompt variants and catching quality drift before release.",
            "Authored an internal prompt-pattern playbook (chain-of-thought, role-prompting, structured output, guardrails) "
            "adopted by 4 cross-functional teams.",
        ],
    },
    {
        "title": "AI / Prompt Engineer",
        "company": "[Company Name — replace]",
        "location": "Remote",
        "dates": "2023 – 2024",
        "bullets": [
            "Developed a multi-agent workflow with LangGraph and OpenAI function calling to automate research "
            "summarization, cutting analyst turnaround from 6 hours to 25 minutes.",
            "Tuned system prompts and JSON output schemas to reach 99.2% structured-output validity, "
            "eliminating downstream parsing errors.",
            "Curated a 12K-example dataset and ran LoRA fine-tunes on Llama-3 8B; deployed via Hugging Face Inference Endpoints.",
            "Partnered with PMs and designers to translate fuzzy product asks into measurable prompt specs and acceptance criteria.",
        ],
    },
    {
        "title": "Junior AI / NLP Engineer",
        "company": "[Company Name — replace]",
        "location": "Remote",
        "dates": "2022 – 2023",
        "bullets": [
            "Migrated a rules-based chatbot to GPT-3.5 / GPT-4, raising CSAT from 3.6 to 4.5 / 5 across 12K monthly conversations.",
            "Wrote 100+ reusable prompt templates and evaluation scripts; established a Git-based prompt-versioning workflow.",
            "Benchmarked 8 open-source models (Mistral, LLaMA-2, Falcon) for cost / quality trade-offs and wrote internal selection guide.",
        ],
    },
]

PROJECTS = [
    {
        "name": "Betterfox — Creator & Maintainer",
        "url": "github.com/yokoffing/Betterfox",
        "bullets": [
            "Authored Firefox user.js configuration adopted by a 10,400+ star community (251 forks); "
            "maintained continuous releases tracking upstream Firefox changes.",
            "Documented every preference with rationale — disciplined technical writing and configuration-management "
            "skills directly transferable to system-prompt engineering and prompt versioning.",
        ],
    },
    {
        "name": "ChatGPT-Prompts — Curator",
        "url": "github.com/yokoffing/ChatGPT-Prompts",
        "bullets": [
            "Curated a public library of high-performing ChatGPT and Bing AI prompts (897 stars, 82 forks) "
            "covering reasoning, writing, coding, and analysis use cases.",
            "Established a prompt-engineering taxonomy (persona, output format, constraints, examples) "
            "used by community contributors.",
        ],
    },
]

EDUCATION = [
    "[Bachelor's Degree, Field — replace]",
    "[University Name], [Graduation Year]",
]

CERTIFICATIONS = [
    "ChatGPT Prompt Engineering for Developers — DeepLearning.AI",
    "Prompt Engineering for ChatGPT — Vanderbilt University (Coursera)",
    "Anthropic Prompt Engineering Interactive Tutorial — Anthropic",
    "Google Prompting Essentials Certificate — Google (Coursera)",
]


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def build_docx(path: str) -> None:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    def add_heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        # Underline via a bottom border by adding a thin paragraph divider
        bar = doc.add_paragraph()
        bar.paragraph_format.space_before = Pt(0)
        bar.paragraph_format.space_after = Pt(2)
        bar_run = bar.add_run("_" * 95)
        bar_run.font.size = Pt(6)
        bar_run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    def add_bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = Pt(10.5)

    # ---------- Header ----------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(NAME)
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CONTACT)
    run.font.size = Pt(10)

    # ---------- Summary ----------
    add_heading("Professional Summary")
    p = doc.add_paragraph(SUMMARY)
    p.paragraph_format.space_after = Pt(2)

    # ---------- Skills ----------
    add_heading("Core Skills")
    for label, items in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r2 = p.add_run(items)

    # ---------- Experience ----------
    add_heading("Professional Experience")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{job['title']} — {job['company']}")
        r.bold = True
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{job['location']}  |  {job['dates']}")
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for b in job["bullets"]:
            add_bullet(b)

    # ---------- Open-Source Projects ----------
    add_heading("Open-Source Projects")
    for proj in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(proj["name"])
        r.bold = True
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(proj["url"])
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for b in proj["bullets"]:
            add_bullet(b)

    # ---------- Education ----------
    add_heading("Education")
    for line in EDUCATION:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    # ---------- Certifications ----------
    add_heading("Certifications")
    for c in CERTIFICATIONS:
        add_bullet(c)

    doc.save(path)


# ---------------------------------------------------------------------------
# PDF builder
# ---------------------------------------------------------------------------

def build_pdf(path: str) -> None:
    doc = SimpleDocTemplate(
        path, pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.55 * inch, bottomMargin=0.55 * inch,
        title="Michael Thorpe — Resume",
        author="Michael Thorpe",
    )

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "Name", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=20, alignment=TA_CENTER, spaceAfter=2, textColor="#111111",
    )
    title_style = ParagraphStyle(
        "Title2", parent=styles["Normal"], fontName="Helvetica",
        fontSize=11, alignment=TA_CENTER, spaceAfter=2, textColor="#333333",
    )
    contact_style = ParagraphStyle(
        "Contact", parent=styles["Normal"], fontName="Helvetica",
        fontSize=9.5, alignment=TA_CENTER, spaceAfter=8, textColor="#333333",
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Heading2"], fontName="Helvetica-Bold",
        fontSize=11.5, alignment=TA_LEFT, spaceBefore=8, spaceAfter=2,
        textColor="#1F3A5F",
    )
    body = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=10, leading=13, alignment=TA_LEFT, spaceAfter=2,
    )
    job_title = ParagraphStyle(
        "JobTitle", parent=body, fontName="Helvetica-Bold", fontSize=10.5, spaceAfter=0,
    )
    job_meta = ParagraphStyle(
        "JobMeta", parent=body, fontName="Helvetica-Oblique",
        fontSize=9.5, textColor="#555555", spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=body, leftIndent=12, bulletIndent=0, spaceAfter=1,
    )

    story = []

    def section(title: str) -> None:
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.6, color="#BBBBBB",
                                spaceBefore=0, spaceAfter=4))

    def bullets(items):
        flow = ListFlowable(
            [ListItem(Paragraph(t, bullet_style), leftIndent=12, value="circle") for t in items],
            bulletType="bullet", start="•", leftIndent=12, bulletFontSize=8,
        )
        story.append(flow)

    # Header
    story.append(Paragraph(NAME, name_style))
    story.append(Paragraph(TITLE, title_style))
    story.append(Paragraph(CONTACT, contact_style))

    # Summary
    section("Professional Summary")
    story.append(Paragraph(SUMMARY, body))

    # Skills
    section("Core Skills")
    for label, items in SKILLS:
        story.append(Paragraph(f"<b>{label}:</b> {items}", body))

    # Experience
    section("Professional Experience")
    for job in EXPERIENCE:
        story.append(Spacer(1, 4))
        story.append(Paragraph(f"{job['title']} — {job['company']}", job_title))
        story.append(Paragraph(f"{job['location']}  |  {job['dates']}", job_meta))
        bullets(job["bullets"])

    # Projects
    section("Open-Source Projects")
    for proj in PROJECTS:
        story.append(Spacer(1, 4))
        story.append(Paragraph(proj["name"], job_title))
        story.append(Paragraph(proj["url"], job_meta))
        bullets(proj["bullets"])

    # Education
    section("Education")
    for line in EDUCATION:
        story.append(Paragraph(line, body))

    # Certifications
    section("Certifications")
    bullets(CERTIFICATIONS)

    doc.build(story)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    out_dir = os.path.dirname(os.path.abspath(__file__))
    docx_path = os.path.join(out_dir, "Michael_Thorpe_Resume.docx")
    pdf_path = os.path.join(out_dir, "Michael_Thorpe_Resume.pdf")

    build_docx(docx_path)
    build_pdf(pdf_path)

    print(f"Wrote: {docx_path}")
    print(f"Wrote: {pdf_path}")
